import streamlit as st
import pandas as pd
import plotly.express as px
import plotly.graph_objects as go
from datetime import datetime, date, timedelta
import uuid
import gspread
from google.oauth2.service_account import Credentials

params = st.query_params
MODO_FERIA = params.get("modo") == "feria"

st.set_page_config(
    page_title="Náutica Viamar Feria" if MODO_FERIA else "Náutica Viamar — NautiCRM",
    page_icon="⚓",
    layout="centered" if MODO_FERIA else "wide",
    initial_sidebar_state="collapsed" if MODO_FERIA else "expanded"
)

# ─── LOGIN ───────────────────────────────────────────────────────────────────
def check_login():
    if st.session_state.get("logged_in"): return True
    st.markdown("""
    <style>
    .block-container{display:flex;justify-content:center}
    </style>
    <div style="max-width:380px;margin:60px auto 0;background:#0d1e35;border:1px solid #1a3050;border-radius:14px;padding:40px 36px;text-align:center">
      <div style="font-family:serif;color:#c9a84c;font-size:1.9rem;margin-bottom:4px">⚓ NautiCRM</div>
      <div style="color:#7a8fa6;font-size:0.82rem;margin-bottom:28px">Gestión de Ventas Náuticas</div>
    </div>
    """, unsafe_allow_html=True)
    with st.form("login_form"):
        usuario  = st.text_input("👤 Usuario",    placeholder="tu usuario")
        password = st.text_input("🔑 Contraseña", type="password", placeholder="••••••••")
        ok = st.form_submit_button("Entrar", use_container_width=True)
        if ok:
            usuarios = dict(st.secrets.get("usuarios", {}))
            if usuario in usuarios and usuarios[usuario] == password:
                st.session_state["logged_in"]       = True
                st.session_state["usuario_activo"]  = usuario
                st.rerun()
            else:
                st.error("Usuario o contraseña incorrectos.")
    return False

if not MODO_FERIA and not check_login():
    st.stop()

STAGES = ["Prospecto","Contactado","Interés Confirmado","Propuesta Enviada","Negociación","Cerrado Ganado","Cerrado Perdido","En Pausa / Recuperable"]
FUNNEL_STAGES = [s for s in STAGES if s not in ["Cerrado Perdido","En Pausa / Recuperable"]]
PAUSA_MESES_ARCHIVO = 6
IDIOMAS = ["Español","Inglés","Francés","Italiano","Alemán","Portugués","Holandés","Ruso","Árabe","Chino","Otro"]
STAGE_COLORS = {
    "Prospecto":"#1a4a8a","Contactado":"#2563eb","Interés Confirmado":"#7c3aed",
    "Propuesta Enviada":"#b8860b","Negociación":"#ea580c","Cerrado Ganado":"#16a34a",
    "Cerrado Perdido":"#dc2626","En Pausa / Recuperable":"#374151",
}

st.markdown('''
<style>
@import url('https://fonts.googleapis.com/css2?family=Playfair+Display:wght@600;700&family=Inter:wght@300;400;500;600;700&display=swap');
html,body,[class*="css"]{font-family:'Inter',sans-serif;background-color:#060e1a!important;color:#e8e0d0!important}
.main{background-color:#060e1a!important}.block-container{padding:1.5rem 2rem!important}
section[data-testid="stSidebar"]{background:#0a1628!important;border-right:1px solid #1a3050}
section[data-testid="stSidebar"] *{color:#e8e0d0!important}
h1,h2,h3{font-family:'Playfair Display',serif!important;color:#c9a84c!important}
[data-testid="metric-container"]{background:#0d1e35;border:1px solid #1a3050;border-radius:10px;padding:1rem 1.2rem}
[data-testid="metric-container"] label{color:#7a8fa6!important;font-size:0.7rem!important;text-transform:uppercase;letter-spacing:1px}
[data-testid="metric-container"] [data-testid="stMetricValue"]{color:#c9a84c!important;font-family:monospace!important;font-size:1.6rem!important}
.stTextInput input,.stNumberInput input,.stTextArea textarea{background:#091220!important;color:#e8e0d0!important;border:1px solid #1a3050!important}
.stButton>button,[data-testid="stFormSubmitButton"]>button{background:#c9a84c!important;color:#0a1628!important;font-weight:700!important;border:none!important;border-radius:6px!important}
.stButton>button:hover,[data-testid="stFormSubmitButton"]>button:hover{opacity:.88!important}
.stDataFrame{border:1px solid #1a3050!important;border-radius:8px!important}
.stTabs [data-baseweb="tab-list"]{background:#0a1628;border-radius:8px;gap:4px;padding:4px}
.stTabs [data-baseweb="tab"]{background:transparent!important;color:#7a8fa6!important;border-radius:6px!important}
.stTabs [aria-selected="true"]{background:#c9a84c!important;color:#0a1628!important;font-weight:700!important}
.streamlit-expanderHeader{background:#0d1e35!important;border:1px solid #1a3050!important;border-radius:8px!important}
hr{border-color:#1a3050!important}
</style>''', unsafe_allow_html=True)

# ─── GOOGLE SHEETS ────────────────────────────────────────────────────────────
@st.cache_resource
def get_client():
    creds = Credentials.from_service_account_info(
        dict(st.secrets["gcp_service_account"]),
        scopes=["https://spreadsheets.google.com/feeds","https://www.googleapis.com/auth/drive"]
    )
    return gspread.authorize(creds)

def get_sheet(tab):
    sh = get_client().open_by_key(st.secrets["spreadsheet_id"])
    try: return sh.worksheet(tab)
    except gspread.WorksheetNotFound:
        ws = sh.add_worksheet(title=tab, rows=1000, cols=30)
        return ws

LEAD_COLS = ["id","nombre","empresa","telefono","email","idioma","tipoEmbarcacion","modeloEslora","presupuesto","usoPrevisto","asignadoA","etapa","probabilidad","valorOperacion","fuenteLead","proximaAccion","fechaProximaAccion","historial","fechaCreacion","ultimaActualizacion"]

@st.cache_data(ttl=30)
def load_leads():
    ws = get_sheet("Leads")
    all_values = ws.get_all_values()
    # Si la hoja está vacía del todo, crear cabeceras
    if not all_values:
        ws.append_row(LEAD_COLS)
        return []
    # Si solo hay una fila (las cabeceras), no hay leads aún
    if len(all_values) <= 1:
        return []
    headers = all_values[0]
    data = [dict(zip(headers, row)) for row in all_values[1:]]
    rows = []
    for r in data:
        # Saltar filas que sean cabeceras duplicadas
        if r.get("id") == "id" or r.get("nombre") == "nombre":
            continue
        lead = {c: r.get(c,"") for c in LEAD_COLS}
        lead["presupuesto"]    = int(float(lead["presupuesto"]))    if lead["presupuesto"]    else 0
        lead["valorOperacion"] = int(float(lead["valorOperacion"])) if lead["valorOperacion"] else 0
        lead["probabilidad"]   = int(float(lead["probabilidad"]))   if lead["probabilidad"]   else 0
        # historial como lista
        h = lead["historial"]
        if isinstance(h, str) and h:
            entries = []
            for e in h.split(";;"):
                parts = e.split("|",2)
                if len(parts)==3: entries.append({"fecha":parts[0],"tipo":parts[1],"nota":parts[2]})
            lead["historial"] = entries
        else:
            lead["historial"] = []
        rows.append(lead)
    return rows

@st.cache_data(ttl=30)
def load_config():
    ws = get_sheet("Config")
    all_values = ws.get_all_values()
    vendedores = ["Vendedor 1","Vendedor 2","Vendedor 3"]
    boat_types = ["Velero","Motor","Catamarán","Zodiac","Charter","Jeanneau","Beneteau","Sunseeker","Princess","Azimut","Ferretti","Bavaria","Hanse","Lagoon","Otro"]
    sources    = ["Feria Náutica","Web","Referido","RRSS","Llamada Fría","Otro"]
    archivo    = []
    if not all_values or len(all_values) <= 1:
        data = []
    else:
        headers = all_values[0]
        data = [dict(zip(headers, row)) for row in all_values[1:]]
    if not data:
        ws.append_row(["key","value"])
        for i,v in enumerate(vendedores): ws.append_row([f"v{i+1}",v])
        ws.append_row(["boat_types", "||".join(boat_types)])
        ws.append_row(["sources", "||".join(sources)])
        return vendedores, boat_types, sources, archivo
    df = pd.DataFrame(data)
    vv = df[df["key"].isin(["v1","v2","v3"])]["value"].tolist()
    if vv: vendedores = vv
    bt_row = df[df["key"]=="boat_types"]["value"].tolist()
    if bt_row: boat_types = bt_row[0].split("||")
    src_row = df[df["key"]=="sources"]["value"].tolist()
    if src_row: sources = src_row[0].split("||")
    arch_row = df[df["key"]=="archivo_frio"]["value"].tolist()
    if arch_row and arch_row[0]:
        import json
        try: archivo = json.loads(arch_row[0])
        except: archivo = []
    return vendedores, boat_types, sources, archivo

def save_lead(lead, is_new=True):
    ws = get_sheet("Leads")
    hist_str = ";;".join([f"{h['fecha']}|{h['tipo']}|{h['nota']}" for h in lead.get("historial",[])])
    row = [lead.get(c,"") for c in LEAD_COLS[:-1]] + [hist_str if c=="historial" else lead.get(c,"") for c in ["ultimaActualizacion"]]
    row = []
    for c in LEAD_COLS:
        if c == "historial": row.append(hist_str)
        else: row.append(lead.get(c,""))
    if is_new:
        ws.append_row(row)
    else:
        all_values = ws.get_all_values()
        if all_values and len(all_values) > 1:
            headers = all_values[0]
            data = [dict(zip(headers, row)) for row in all_values[1:]]
            for i,r in enumerate(data):
                if r.get("id") == lead["id"]:
                    ws.update(f"A{i+2}:{chr(64+len(LEAD_COLS))}{i+2}", [row])
                    break
    load_leads.clear()

def delete_lead(lead_id):
    ws = get_sheet("Leads")
    all_values = ws.get_all_values()
    if all_values and len(all_values) > 1:
        headers = all_values[0]
        data = [dict(zip(headers, row)) for row in all_values[1:]]
        for i,r in enumerate(data):
            if r.get("id") == lead_id:
                ws.delete_rows(i+2); break
    load_leads.clear()

def save_config(vendedores, boat_types, sources, archivo):
    import json
    ws = get_sheet("Config")
    ws.clear()
    rows = [
        ["key", "value"],
        ["v1", vendedores[0]],
        ["v2", vendedores[1]],
        ["v3", vendedores[2]],
        ["boat_types", "||".join(boat_types)],
        ["sources", "||".join(sources)],
        ["archivo_frio", json.dumps(archivo, ensure_ascii=False)],
    ]
    ws.update("A1", rows)
    load_config.clear()

# ─── HELPERS ──────────────────────────────────────────────────────────────────
def fmt_eur(n):
    try: return f"€{int(n):,}".replace(",",".")
    except: return "—"

def urg_emoji(ds):
    if not ds: return "⚪"
    try:
        diff=(datetime.strptime(str(ds),"%Y-%m-%d").date()-date.today()).days
        return "🔴" if diff<0 else "🟡" if diff==0 else "🟢"
    except: return "⚪"

def days_since(ds):
    try: return (date.today()-datetime.strptime(str(ds),"%Y-%m-%d").date()).days
    except: return 0

def months_since(ds):
    if not ds: return 999
    try:
        d=datetime.strptime(str(ds),"%Y-%m-%d").date()
        return (date.today().year-d.year)*12+(date.today().month-d.month)
    except: return 0

# ─── CARGAR DATOS ─────────────────────────────────────────────────────────────
try:
    all_leads_raw = load_leads()
    vendedores, boat_types, sources, archivo_frio = load_config()
except Exception as e:
    st.error(f"❌ Error conectando con Google Sheets: {e}")
    st.info("Comprueba que los Secrets están bien configurados en Streamlit Cloud.")
    st.stop()

# Auto-archivar
def auto_archivar(leads):
    activos, archivados = [], []
    for l in leads:
        if l.get("etapa")=="En Pausa / Recuperable" and months_since(l.get("ultimaActualizacion") or l.get("fechaCreacion",""))>=PAUSA_MESES_ARCHIVO:
            archivados.append({**l,"fechaArchivo":str(date.today())}); continue
        activos.append(l)
    return activos, archivados

activos, nuevos_arch = auto_archivar(all_leads_raw)
if nuevos_arch:
    archivo_frio.extend(nuevos_arch)
    save_config(vendedores, boat_types, archivo_frio)
    for l in nuevos_arch: delete_lead(l["id"])
    all_leads_raw = activos

# ══ MODO FERIA ════════════════════════════════════════════════════════════════
if MODO_FERIA:
    st.markdown("""<style>
    .block-container{padding:1rem!important}
    .stButton>button{width:100%;padding:.8rem!important;font-size:1rem!important;border-radius:10px!important;margin-top:8px!important}
    </style>""", unsafe_allow_html=True)
    st.markdown("<h1 style='text-align:center;font-size:1.5rem;margin-bottom:0'>⚓ Alta Rápida</h1>", unsafe_allow_html=True)
    st.markdown("<p style='text-align:center;color:#7a8fa6;font-size:0.82rem;margin-top:2px'>Modo Feria — registro rápido de contactos</p>", unsafe_allow_html=True)
    st.markdown("---")
    with st.form("feria_form", clear_on_submit=True):
        nombre   = st.text_input("👤 Nombre completo *")
        empresa  = st.text_input("🏢 Empresa")
        telefono = st.text_input("📱 Teléfono / WhatsApp")
        email    = st.text_input("📧 Email")
        idioma   = st.selectbox("🌍 Idioma", IDIOMAS)
        tipo     = st.selectbox("🚢 Tipo / Marca", boat_types)
        modelo   = st.text_input("📐 Modelo / Eslora")
        presu    = st.number_input("💶 Presupuesto aprox. (€)", min_value=0, step=10000)
        asig     = st.selectbox("👔 Asignar a", vendedores)
        nota_ini = st.text_area("📝 Nota inicial", placeholder="Observaciones...", height=80)
        ok = st.form_submit_button("✅ GUARDAR LEAD", use_container_width=True)
        if ok:
            if not nombre.strip(): st.error("El nombre es obligatorio.")
            else:
                nuevo = {"id":str(uuid.uuid4()),"nombre":nombre,"empresa":empresa,"telefono":telefono,"email":email,"idioma":idioma,"tipoEmbarcacion":tipo,"modeloEslora":modelo,"presupuesto":int(presu),"usoPrevisto":"","asignadoA":asig,"etapa":"Prospecto","probabilidad":10,"valorOperacion":int(presu),"fuenteLead":"Feria Náutica","proximaAccion":"Contactar post-feria","fechaProximaAccion":str(date.today()+timedelta(days=3)),"historial":[{"fecha":str(date.today()),"tipo":"Nota","nota":nota_ini.strip()}] if nota_ini.strip() else [],"fechaCreacion":str(date.today()),"ultimaActualizacion":str(date.today())}
                save_lead(nuevo, is_new=True)
                st.success(f"✅ **{nombre}** guardado."); st.balloons()
    st.markdown("---")
    recientes=[l for l in all_leads_raw if l.get("fuenteLead")=="Feria Náutica"][-5:]
    if recientes:
        st.caption(f"Últimos {len(recientes)} leads de feria:")
        for l in reversed(recientes):
            st.markdown(f"<div style='background:#0d1e35;border:1px solid #1a3050;border-radius:8px;padding:10px 14px;margin-bottom:6px'><span style='color:#c9a84c;font-weight:700'>{l['nombre']}</span><span style='color:#7a8fa6;font-size:0.78rem'> · {l.get('empresa','')} · {l.get('tipoEmbarcacion','')}</span></div>", unsafe_allow_html=True)
    st.stop()

# ══ SIDEBAR ═══════════════════════════════════════════════════════════════════
with st.sidebar:
    try:
        st.image("LOGO-viamar - bueno.jpg", use_container_width=True)
    except:
        pass
    st.markdown("## ⚓ NautiCRM")
    st.markdown("*Náutica Viamar*")
    st.markdown("---")
    active_user  = st.selectbox("👤 Usuario activo", vendedores)
    my_portfolio = st.checkbox("📂 Ver solo mi cartera")
    st.markdown("---")
    # Navegación con soporte de redirección desde kanban/lista
    if "nav_page" not in st.session_state:
        st.session_state["nav_page"] = "⊞ Funnel Kanban"
    if st.session_state.get("_nav_request"):
        st.session_state["nav_page"] = st.session_state["_nav_request"]
        st.session_state["_nav_request"] = None
    if st.session_state.get("goto_lead"):
        st.session_state["nav_page"] = "➕ Nuevo / Editar Lead"
    page = st.radio("Navegación",[
        "⊞ Funnel Kanban","≡ Lista de Leads","📊 Informes",
        "➕ Nuevo / Editar Lead","📅 Próximas Acciones","🧊 Archivo Frío","⚙️ Configuración"
    ], key="nav_page")
    if page != "➕ Nuevo / Editar Lead":
        st.session_state["goto_lead"] = None
    st.markdown("---")
    if st.button("🔄 Actualizar datos"): load_leads.clear(); load_config.clear(); st.rerun()
    n_arch=len(archivo_frio)
    if n_arch>0: st.markdown(f"<div style='background:#1a3050;border-radius:6px;padding:6px 10px;font-size:0.75rem;color:#7a8fa6'>🧊 Archivo Frío: <b style='color:#e8e0d0'>{n_arch}</b></div>", unsafe_allow_html=True)
    st.markdown("<br>", unsafe_allow_html=True)
    app_url = st.secrets.get("app_url","https://nauticrm-trn2jrtqldn9vfd4zic6hz.streamlit.app")
    st.markdown(f"<a href='{app_url}/?modo=feria' target='_blank' style='background:#c9a84c;color:#0a1628;padding:8px 14px;border-radius:6px;font-weight:700;font-size:0.8rem;text-decoration:none'>📱 Abrir modo Feria</a>", unsafe_allow_html=True)

leads = all_leads_raw
if my_portfolio: leads=[l for l in leads if l["asignadoA"]==active_user]

# ══ KANBAN ════════════════════════════════════════════════════════════════════
if "Kanban" in page:
    st.markdown("## ⊞ Funnel de Ventas")
    st.markdown("<br>", unsafe_allow_html=True)
    funnel_leads=[l for l in leads if l["etapa"] in FUNNEL_STAGES]
    cols=st.columns(len(FUNNEL_STAGES))
    for col,stage in zip(cols,FUNNEL_STAGES):
        cards=[l for l in funnel_leads if l["etapa"]==stage]
        total=sum(l.get("valorOperacion",0) for l in cards)
        color=STAGE_COLORS[stage]
        with col:
            st.markdown(f"""<div style="background:{color};border-radius:6px 6px 0 0;padding:5px 10px;display:flex;justify-content:space-between;align-items:center">
                <span style="color:white;font-size:0.6rem;font-weight:700;text-transform:uppercase;white-space:nowrap;overflow:hidden;text-overflow:ellipsis;max-width:75%">{stage}</span>
                <span style="background:rgba(0,0,0,0.35);color:white;border-radius:10px;padding:1px 8px;font-size:0.85rem;font-weight:700;flex-shrink:0">{len(cards)}</span>
            </div><div style="background:#091220;border-radius:0 0 6px 6px;padding:5px;min-height:80px">
            {"<div style='text-align:center;padding:4px 2px 6px;font-family:monospace;font-size:1rem;font-weight:700;color:#c9a84c'>"+fmt_eur(total)+"</div>" if total>0 else ""}""", unsafe_allow_html=True)
            for l in cards:
                d=days_since(l.get("ultimaActualizacion") or l.get("fechaCreacion",""))
                btn_label=f"{urg_emoji(l.get('fechaProximaAccion'))} {l['nombre']}|{l.get('modeloEslora') or l['tipoEmbarcacion']}|{fmt_eur(l.get('valorOperacion',0))} · {l['asignadoA'].replace('Vendedor','V.')} · {d}d"
                st.markdown(f"""<style>
                div[data-testid="stButton"] button[kind="secondary"]{{
                    background:#0d1e35!important;border:1px solid #1a3050!important;
                    border-left:3px solid {color}!important;border-radius:6px!important;
                    padding:7px 9px!important;margin:2px 0!important;width:100%!important;
                    text-align:left!important;color:#e8e0d0!important;font-size:0.72rem!important;
                    white-space:normal!important;line-height:1.4!important;
                }}
                div[data-testid="stButton"] button[kind="secondary"]:hover{{
                    background:#1a3050!important;border-color:{color}!important;
                }}
                </style>""", unsafe_allow_html=True)
                if st.button(btn_label, key=f"kb_{l['id']}"):
                    st.session_state["goto_lead"] = l["nombre"]; st.rerun()
            st.markdown("</div>", unsafe_allow_html=True)

    st.markdown("<br>", unsafe_allow_html=True)
    with st.expander("⚡ Cambio rápido de estado"):
        c1,c2,c3=st.columns(3)
        nombres_k=[l["nombre"] for l in sorted(all_leads_raw,key=lambda l:l.get("fechaCreacion",""),reverse=True) if l["etapa"] not in ["Cerrado Ganado","Cerrado Perdido"]]
        sel_k=c1.selectbox("Lead",["— Selecciona —"]+nombres_k,key="sel_k")
        nueva_etapa=c2.selectbox("Nueva etapa",STAGES,key="etapa_k")
        if c3.button("✅ Cambiar"):
            if sel_k!="— Selecciona —":
                lead_upd=next(l for l in all_leads_raw if l["nombre"]==sel_k)
                lead_upd["etapa"]=nueva_etapa; lead_upd["ultimaActualizacion"]=str(date.today())
                save_lead(lead_upd, is_new=False)
                st.success(f"✅ {sel_k} → {nueva_etapa}"); st.rerun()

    st.markdown("<div style='margin:8px 0 16px;padding:8px 14px;background:#0d1e35;border:1px solid #1a3050;border-radius:8px;font-size:0.78rem;color:#7a8fa6'>🔴 Vencida &nbsp;|&nbsp; 🟡 Hoy &nbsp;|&nbsp; 🟢 Futura &nbsp;|&nbsp; ⚪ Sin fecha asignada</div>", unsafe_allow_html=True)
    st.markdown("<hr style='border-color:#1a3050'>", unsafe_allow_html=True)
    bc1,bc2,bc3=st.columns(3)
    for col,stage in zip([bc1,bc2,bc3],["Cerrado Ganado","Cerrado Perdido","En Pausa / Recuperable"]):
        # Excluir cerrados/pausa con más de 13 meses sin actividad
        cards=[l for l in all_leads_raw if l["etapa"]==stage and months_since(l.get("ultimaActualizacion") or l.get("fechaCreacion",""))<=13]
        total=sum(l.get("valorOperacion",0) for l in cards)
        color=STAGE_COLORS[stage]
        icon="✅" if stage=="Cerrado Ganado" else ("❌" if stage=="Cerrado Perdido" else "⏸️")
        with col:
            st.markdown(f"""<div style="background:{color};border-radius:6px 6px 0 0;padding:5px 12px;display:flex;justify-content:space-between;align-items:center">
                <span style="color:white;font-size:0.65rem;font-weight:700;text-transform:uppercase">{icon} {stage}</span>
                <span style="background:rgba(0,0,0,0.3);color:white;border-radius:10px;padding:0 6px;font-size:0.85rem;font-weight:700">{len(cards)}</span>
            </div><div style="background:#091220;border-radius:0 0 6px 6px;padding:5px;min-height:80px">
            {"<div style='text-align:center;padding:4px 2px 6px;font-family:monospace;font-size:1rem;font-weight:700;color:#c9a84c'>"+fmt_eur(total)+"</div>" if total>0 else ""}""", unsafe_allow_html=True)
            for l in cards[:5]:
                st.markdown(f"""<div style="background:#0d1e35;border:1px solid #1a3050;border-left:3px solid {color};border-radius:5px;padding:5px 8px;margin:3px 0;font-size:0.7rem">
                    <div style="font-weight:600;color:#e8e0d0;white-space:nowrap;overflow:hidden;text-overflow:ellipsis">{l['nombre']}</div>
                    <div style="color:#7a8fa6;font-size:0.62rem">{l['empresa']}·{fmt_eur(l.get('valorOperacion',0))}</div></div>""", unsafe_allow_html=True)
            if len(cards)>5: st.markdown(f"<small style='color:#7a8fa6;padding:0 5px'>+{len(cards)-5} más</small>", unsafe_allow_html=True)
            st.markdown("</div>", unsafe_allow_html=True)

# ══ LISTA ══════════════════════════════════════════════════════════════════════
elif "Lista" in page:
    st.markdown("## ≡ Lista de Leads")
    c1,c2,c3,c4=st.columns(4)
    q=c1.text_input("🔍 Buscar",placeholder="Nombre, empresa...")
    fv=c2.selectbox("Vendedor",["Todos"]+vendedores)
    fe=c3.selectbox("Etapa",["Todas"]+STAGES)
    ft=c4.selectbox("Tipo",["Todos"]+boat_types)
    filtered=leads
    if q: filtered=[l for l in filtered if q.lower() in (l.get("nombre","")+" "+l.get("empresa","")+" "+l.get("modeloEslora","")).lower()]
    if fv!="Todos": filtered=[l for l in filtered if l["asignadoA"]==fv]
    if fe!="Todas": filtered=[l for l in filtered if l["etapa"]==fe]
    if ft!="Todos": filtered=[l for l in filtered if l["tipoEmbarcacion"]==ft]
    st.caption(f"{len(filtered)} leads")
    if not filtered: st.info("Sin resultados.")
    else:
        for l in filtered:
            c1,c2=st.columns([8,1])
            with c1:
                color=STAGE_COLORS.get(l["etapa"],"#1a3050")
                st.markdown(f"""<div style="background:#0d1e35;border:1px solid #1a3050;border-left:3px solid {color};border-radius:6px;padding:8px 12px;margin-bottom:4px;display:flex;justify-content:space-between;align-items:center">
                    <div><span style="color:#e8e0d0;font-weight:700;font-size:0.85rem">{urg_emoji(l.get('fechaProximaAccion'))} {l['nombre']}</span>
                    <span style="background:{color};color:white;border-radius:10px;padding:1px 7px;font-size:0.62rem;font-weight:700;margin-left:8px">{l['etapa']}</span>
                    <span style="color:#7a8fa6;font-size:0.72rem;margin-left:8px">{l.get('empresa','')}</span></div>
                    <div style="text-align:right"><span style="color:#2ecc71;font-family:monospace;font-size:0.78rem;font-weight:700">{fmt_eur(l.get('valorOperacion',0))}</span>
                    <span style="color:#7a8fa6;font-size:0.7rem;margin-left:8px">{l.get('tipoEmbarcacion','')} {l.get('modeloEslora','')}</span></div>
                </div>""", unsafe_allow_html=True)
            with c2:
                if st.button("✏️", key=f"lst_{l['id']}", help="Ir a ficha"):
                    st.session_state["goto_lead"] = l["nombre"]; st.rerun()
        st.markdown("<br>", unsafe_allow_html=True)
        rows_exp=[{"Urg":urg_emoji(l.get("fechaProximaAccion")),"Nombre":l["nombre"],"Empresa":l["empresa"],"Idioma":l.get("idioma","—"),"Embarcación":f"{l['tipoEmbarcacion']}·{l['modeloEslora']}","Etapa":l["etapa"],"Valor €":fmt_eur(l.get("valorOperacion",0)),"Prob%":f"{l.get('probabilidad',0)}%","Vendedor":l["asignadoA"],"Próx. Acción":l.get("fechaProximaAccion","—")} for l in filtered]
        import io as _io
        _buf_lista=_io.BytesIO()
        pd.DataFrame(rows_exp).to_excel(_buf_lista,index=False,engine="openpyxl")
        st.download_button("⬇️ Exportar Excel",data=_buf_lista.getvalue(),file_name="nauticrm_leads.xlsx",mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")

# ══ INFORMES ══════════════════════════════════════════════════════════════════
elif "Informes" in page:
    st.markdown("## 📊 Informes y Análisis")
    active=[l for l in all_leads_raw if l["etapa"] not in ["Cerrado Ganado","Cerrado Perdido","En Pausa / Recuperable"]]
    pipeline=sum(l.get("valorOperacion",0) for l in active)
    forecast=sum(l.get("valorOperacion",0)*l.get("probabilidad",0)/100 for l in all_leads_raw if l["etapa"]!="Cerrado Perdido")
    no_act=[l for l in active if days_since(l.get("ultimaActualizacion") or l.get("fechaCreacion",""))>7]
    k1,k2,k3,k4=st.columns(4)
    k1.metric("Total Leads",len(all_leads_raw),f"{len(active)} activos")
    k2.metric("Pipeline Activo",fmt_eur(pipeline))
    k3.metric("Forecast Mensual",fmt_eur(forecast),"Valor × Prob%")
    k4.metric("Sin actividad +7d",len(no_act),delta="⚠️ Revisar" if no_act else "✅ OK",delta_color="off")
    st.markdown("<br>", unsafe_allow_html=True)
    tab1,tab2,tab3,tab4=st.tabs(["🔻 Embudo","📊 Pipeline & Fuentes","📋 Actividad","🤖 Informe IA"])
    with tab1:
        st.markdown("#### Embudo de Ventas")
        # Todas las etapas: funnel normal + cerrados/pausa con <13 meses
        ALL_FUNNEL = ["Prospecto","Contactado","Interés Confirmado","Propuesta Enviada","Negociación","Cerrado Ganado","Cerrado Perdido","En Pausa / Recuperable"]
        def en_funnel(l):
            if l["etapa"] in FUNNEL_STAGES: return True
            if l["etapa"] in ["Cerrado Ganado","Cerrado Perdido","En Pausa / Recuperable"]:
                return months_since(l.get("ultimaActualizacion") or l.get("fechaCreacion","")) <= 13
            return False
        leads_funnel = [l for l in all_leads_raw if en_funnel(l)]
        fd=[{"Etapa":s,"Leads":len([l for l in leads_funnel if l["etapa"]==s]),"Valor":sum(l.get("valorOperacion",0) for l in leads_funnel if l["etapa"]==s),"Color":STAGE_COLORS[s]} for s in ALL_FUNNEL]
        fd=[d for d in fd if d["Leads"]>0 or d["Valor"]>0]
        if not fd:
            st.info("Sin datos de pipeline todavía.")
        else:
            c1,c2=st.columns(2)
            with c1:
                st.markdown("**Por número de leads**")
                fig1=go.Figure(go.Bar(
                    y=[d["Etapa"] for d in fd],
                    x=[1]*len(fd),
                    orientation="h",
                    text=[f"<b>{d['Leads']}</b> leads" for d in fd],
                    textposition="inside",
                    textfont=dict(color="white",size=15,family="Arial Bold"),
                    marker=dict(color=[d["Color"] for d in fd],line=dict(color="rgba(0,0,0,0.15)",width=1)),
                    insidetextanchor="middle",
                ))
                fig1.update_layout(
                    plot_bgcolor="#091220",paper_bgcolor="#0d1e35",font_color="#e8e0d0",
                    margin=dict(t=20,b=20,l=10,r=10),height=420,
                    xaxis=dict(showgrid=False,showticklabels=False,range=[0,1]),
                    yaxis=dict(tickfont=dict(size=11,color="white"),autorange="reversed"),
                    bargap=0.25,
                    uniformtext=dict(mode="hide",minsize=10),
                )
                fig1.update_traces(width=0.6)
                st.plotly_chart(fig1,use_container_width=True)
            with c2:
                st.markdown("**Por valor (€)**")
                fig2=go.Figure(go.Bar(
                    y=[d["Etapa"] for d in fd],
                    x=[1]*len(fd),
                    orientation="h",
                    text=[f"<b>{fmt_eur(d['Valor'])}</b>" for d in fd],
                    textposition="inside",
                    textfont=dict(color="white",size=16,family="Arial Bold"),
                    marker=dict(color=[d["Color"] for d in fd],line=dict(color="rgba(0,0,0,0.15)",width=1)),
                    insidetextanchor="middle",
                ))
                fig2.update_layout(
                    plot_bgcolor="#091220",paper_bgcolor="#0d1e35",font_color="#e8e0d0",
                    margin=dict(t=20,b=20,l=10,r=10),height=420,
                    xaxis=dict(showgrid=False,showticklabels=False,range=[0,1]),
                    yaxis=dict(tickfont=dict(size=11,color="white"),autorange="reversed"),
                    bargap=0.25,
                    uniformtext=dict(mode="hide",minsize=10),
                )
                fig2.update_traces(width=0.6)
                st.plotly_chart(fig2,use_container_width=True)
            rows_f=[{"Etapa":d["Etapa"],"Leads":d["Leads"],"Valor Pipeline":fmt_eur(d["Valor"]),"Ticket Medio":fmt_eur(d["Valor"]/d["Leads"]) if d["Leads"]>0 else "—"} for d in fd]
            st.dataframe(pd.DataFrame(rows_f),use_container_width=True,hide_index=True)
    with tab2:
        c1,c2=st.columns(2)
        with c1:
            st.markdown("#### Pipeline por Vendedor")
            pv={v:sum(l.get("valorOperacion",0) for l in active if l["asignadoA"]==v) for v in vendedores}
            df_pv=pd.DataFrame({"Vendedor":list(pv.keys()),"Valor":list(pv.values())})
            fig=px.bar(df_pv,x="Vendedor",y="Valor",color_discrete_sequence=["#c9a84c"],text=df_pv["Valor"].apply(fmt_eur))
            fig.update_layout(plot_bgcolor="#091220",paper_bgcolor="#0d1e35",font_color="#e8e0d0",showlegend=False,xaxis=dict(gridcolor="#1a3050"),yaxis=dict(gridcolor="#1a3050"),margin=dict(t=20,b=10,l=10,r=10))
            fig.update_traces(textposition="outside",textfont_color="white"); st.plotly_chart(fig,use_container_width=True)
        with c2:
            st.markdown("#### Leads por Fuente")
            src={}
            for l in all_leads_raw: src[l.get("fuenteLead","Otro")]=src.get(l.get("fuenteLead","Otro"),0)+1
            fig2=px.pie(pd.DataFrame({"Fuente":list(src.keys()),"N":list(src.values())}),names="Fuente",values="N",hole=0.52,color_discrete_sequence=["#c9a84c","#2563eb","#7c3aed","#ea580c","#16a34a","#dc2626"])
            fig2.update_layout(plot_bgcolor="#091220",paper_bgcolor="#0d1e35",font_color="white",margin=dict(t=20,b=10,l=10,r=10),legend=dict(font=dict(color="white")))
            fig2.update_traces(textfont_color="white",textfont_size=13)
            st.plotly_chart(fig2,use_container_width=True)
        st.markdown("#### Leads por Idioma")
        ic={}
        for l in all_leads_raw: ic[l.get("idioma","—")]=ic.get(l.get("idioma","—"),0)+1
        df_id=pd.DataFrame({"Idioma":list(ic.keys()),"Leads":list(ic.values())}).sort_values("Leads",ascending=False)
        fig3=px.bar(df_id,x="Idioma",y="Leads",color_discrete_sequence=["#2563eb"],text="Leads")
        fig3.update_layout(plot_bgcolor="#091220",paper_bgcolor="#0d1e35",font_color="#e8e0d0",showlegend=False,xaxis=dict(gridcolor="#1a3050"),yaxis=dict(gridcolor="#1a3050"),margin=dict(t=10,b=10,l=10,r=10),height=220)
        fig3.update_traces(textposition="outside",textfont_color="white"); st.plotly_chart(fig3,use_container_width=True)
    with tab3:
        c1,c2=st.columns(2)
        with c1:
            st.markdown("#### Actividad Reciente")
            hist_all=[]
            for l in all_leads_raw:
                for h in l.get("historial",[]): hist_all.append({**h,"lead":l["nombre"]})
            hist_all.sort(key=lambda x:x["fecha"],reverse=True)
            for h in hist_all[:10]:
                st.markdown(f"<div style='padding:7px 0;border-bottom:1px solid #1a3050'><span style='color:#c9a84c;font-size:0.78rem;font-weight:700'>{h['lead']}</span><span style='color:#7a8fa6;font-size:0.72rem'> · {h['fecha']}</span><span style='background:#1a3050;color:#7a8fa6;border-radius:4px;padding:1px 5px;font-size:0.65rem;margin-left:5px'>{h['tipo']}</span><div style='color:#e8e0d0;font-size:0.78rem;margin-top:3px'>{h['nota']}</div></div>", unsafe_allow_html=True)
        with c2:
            st.markdown("#### Sin actividad +7 días")
            if not no_act: st.success("✅ Todo al día.")
            else:
                for l in no_act:
                    d=days_since(l.get("ultimaActualizacion","") or l.get("fechaCreacion",""))
                    st.markdown(f"<div style='display:flex;justify-content:space-between;padding:6px 10px;margin-bottom:5px;background:#0d1e35;border:1px solid #e74c3c33;border-radius:6px'><span style='color:#e8e0d0;font-size:0.82rem'>{l['nombre']}</span><span style='color:#e74c3c;font-size:0.78rem'>{d}d·{l['asignadoA'].replace('Vendedor','V.')}</span></div>", unsafe_allow_html=True)
    with tab4:
        st.markdown("#### 🤖 Informe IA — Análisis y Patrones")
        st.markdown("<div style='color:#7a8fa6;font-size:0.82rem;margin-bottom:16px'>Claude analiza todos los datos del CRM y detecta patrones, tendencias y oportunidades de mejora.</div>",unsafe_allow_html=True)

        # ── Filtro de período ─────────────────────────────────────────────────
        _ia1,_ia2,_ia3=st.columns(3)
        _periodo=_ia1.selectbox("Período",["Todo","Este año","Este trimestre","Este mes","Fechas personalizadas"],key="ia_periodo")
        _hoy=date.today()
        if _periodo=="Este mes":
            _d_ini=_hoy.replace(day=1); _d_fin=_hoy
        elif _periodo=="Este trimestre":
            _q=(_hoy.month-1)//3; _d_ini=date(_hoy.year,_q*3+1,1); _d_fin=_hoy
        elif _periodo=="Este año":
            _d_ini=date(_hoy.year,1,1); _d_fin=_hoy
        elif _periodo=="Fechas personalizadas":
            _d_ini=_ia2.date_input("Desde",value=date(_hoy.year,1,1),key="ia_desde")
            _d_fin=_ia3.date_input("Hasta",value=_hoy,key="ia_hasta")
        else:
            _d_ini=None; _d_fin=None

        def _en_periodo(l):
            if not _d_ini: return True
            try: return _d_ini<=datetime.strptime(str(l.get("fechaCreacion","")),"%Y-%m-%d").date()<=_d_fin
            except: return False

        _leads_ia=[l for l in all_leads_raw if _en_periodo(l)]
        _arch_ia=[l for l in archivo_frio if not _d_ini or (lambda d: _d_ini<=d<=_d_fin if d else False)(
            (lambda s: datetime.strptime(s,"%Y-%m-%d").date() if s else None)(l.get("fechaArchivo","")))]

        st.caption(f"{len(_leads_ia)} leads activos + {len(_arch_ia)} en archivo frío en el período seleccionado")

        if st.button("🤖 Generar informe IA",use_container_width=True):
            # ── Preparar resumen de datos ─────────────────────────────────────
            _todos=_leads_ia+_arch_ia
            _por_etapa={s:len([l for l in _leads_ia if l["etapa"]==s]) for s in STAGES}
            _por_tipo={}
            for l in _todos:
                t=l.get("tipoEmbarcacion","—"); _por_tipo[t]=_por_tipo.get(t,0)+1
            _por_fuente={}
            for l in _todos:
                f=l.get("fuenteLead","—"); _por_fuente[f]=_por_fuente.get(f,0)+1
            _por_idioma={}
            for l in _todos:
                i=l.get("idioma","—"); _por_idioma[i]=_por_idioma.get(i,0)+1
            _actividad=[]
            for l in _leads_ia:
                for h in l.get("historial",[]):
                    _actividad.append({"lead":l["nombre"],"fecha":h["fecha"],"tipo":h["tipo"],"nota":h["nota"][:120]})
            _actividad.sort(key=lambda x:x["fecha"],reverse=True)
            _act_por_tipo={}
            for a in _actividad:
                _act_por_tipo[a["tipo"]]=_act_por_tipo.get(a["tipo"],0)+1
            # Actividad por mes
            _act_por_mes={}
            for a in _actividad:
                try: _mes=a["fecha"][:7]
                except: _mes="—"
                _act_por_mes[_mes]=_act_por_mes.get(_mes,0)+1
            # Altas por mes
            _altas_por_mes={}
            for l in _leads_ia:
                try: _mes=l.get("fechaCreacion","")[:7]
                except: _mes="—"
                if _mes: _altas_por_mes[_mes]=_altas_por_mes.get(_mes,0)+1
            _sin_actividad=[l["nombre"] for l in _leads_ia if not l.get("historial") and l["etapa"] not in ["Cerrado Ganado","Cerrado Perdido"]]
            _estancados=[l["nombre"] for l in _leads_ia if days_since(l.get("ultimaActualizacion") or l.get("fechaCreacion",""))>30 and l["etapa"] not in ["Cerrado Ganado","Cerrado Perdido","En Pausa / Recuperable"]]
            _pipeline_activo=sum(l.get("valorOperacion",0) for l in _leads_ia if l["etapa"] not in ["Cerrado Ganado","Cerrado Perdido","En Pausa / Recuperable"])
            _valor_ganado=sum(l.get("valorOperacion",0) for l in _leads_ia if l["etapa"]=="Cerrado Ganado")
            _valor_perdido=sum(l.get("valorOperacion",0) for l in _leads_ia if l["etapa"]=="Cerrado Perdido")

            _prompt=f"""Eres un consultor experto en ventas náuticas analizando los datos del CRM de Náutica Viamar, una empresa especializada en embarcaciones de recreo.

PERÍODO ANALIZADO: {"Todo el histórico" if not _d_ini else f"{_d_ini} a {_d_fin}"}

## DATOS DEL CRM

### Resumen general
- Total leads en período: {len(_leads_ia)} activos + {len(_arch_ia)} en archivo frío
- Pipeline activo: €{_pipeline_activo:,}
- Valor cerrado ganado: €{_valor_ganado:,}
- Valor cerrado perdido: €{_valor_perdido:,}
- Tasa de conversión: {round(len([l for l in _leads_ia if l['etapa']=='Cerrado Ganado'])/max(len(_leads_ia),1)*100,1)}%

### Leads por etapa
{chr(10).join(f"- {k}: {v}" for k,v in _por_etapa.items() if v>0)}

### Leads por tipo de embarcación
{chr(10).join(f"- {k}: {v}" for k,v in sorted(_por_tipo.items(),key=lambda x:-x[1]))}

### Leads por fuente de captación
{chr(10).join(f"- {k}: {v}" for k,v in sorted(_por_fuente.items(),key=lambda x:-x[1]))}

### Leads por idioma
{chr(10).join(f"- {k}: {v}" for k,v in sorted(_por_idioma.items(),key=lambda x:-x[1]))}

### Actividad registrada ({len(_actividad)} interacciones)
Por tipo: {', '.join(f"{k}: {v}" for k,v in _act_por_tipo.items())}
Por mes: {', '.join(f"{k}: {v}" for k,v in sorted(_act_por_mes.items()))}

### Altas de nuevos leads por mes
{', '.join(f"{k}: {v}" for k,v in sorted(_altas_por_mes.items()))}

### Alertas
- Leads sin ninguna actividad registrada ({len(_sin_actividad)}): {', '.join(_sin_actividad[:10])}{'...' if len(_sin_actividad)>10 else ''}
- Leads estancados +30 días sin actividad ({len(_estancados)}): {', '.join(_estancados[:10])}{'...' if len(_estancados)>10 else ''}

### Muestra de actividad reciente (últimas 20 interacciones)
{chr(10).join(f"- [{a['fecha']}] {a['lead']} | {a['tipo']}: {a['nota']}" for a in _actividad[:20])}

## INSTRUCCIONES
Genera un informe completo en español con estas secciones:
1. **Resumen ejecutivo** (3-4 frases clave)
2. **Análisis de pipeline** (estado del embudo, conversión, valor en juego)
3. **Patrones detectados** (estacionalidad, tipos de cliente que cierran más, fuentes más rentables, idiomas relevantes)
4. **Alertas y riesgos** (leads en riesgo, estancamientos, oportunidades perdidas)
5. **Recomendaciones concretas** (mínimo 5 acciones específicas y prioritarias)

Sé directo, usa datos concretos del informe y enfócate en lo accionable. Formato markdown."""

            with st.spinner("🤖 Claude está analizando los datos..."):
                try:
                    import anthropic as _anthropic
                    _client=_anthropic.Anthropic(api_key=st.secrets["anthropic_api_key"])
                    _msg=_client.messages.create(
                        model="claude-opus-4-5",
                        max_tokens=2000,
                        messages=[{"role":"user","content":_prompt}]
                    )
                    _informe=_msg.content[0].text
                    st.markdown("---")
                    st.markdown(_informe)
                    st.markdown("---")
                    import io as _io_ia, re as _re

                    def _to_pdf(txt):
                        import os as _os
                        from reportlab.lib.pagesizes import A4
                        from reportlab.lib import colors as _rc
                        from reportlab.lib.units import cm
                        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image as _RLImage, HRFlowable
                        from reportlab.lib.styles import ParagraphStyle
                        _buf=_io_ia.BytesIO()
                        _doc=SimpleDocTemplate(_buf,pagesize=A4,leftMargin=2*cm,rightMargin=2*cm,topMargin=2*cm,bottomMargin=2*cm)
                        _GOLD=_rc.HexColor("#C9A84C"); _BLUE=_rc.HexColor("#2563eb"); _DARK=_rc.HexColor("#111111")
                        _s_tit=ParagraphStyle("t",fontName="Helvetica-Bold",fontSize=17,textColor=_GOLD,spaceAfter=4)
                        _s_sub=ParagraphStyle("s",fontName="Helvetica",fontSize=9,textColor=_rc.HexColor("#555555"),spaceAfter=10)
                        _s_h1=ParagraphStyle("h1",fontName="Helvetica-Bold",fontSize=13,textColor=_GOLD,spaceBefore=14,spaceAfter=5)
                        _s_h2=ParagraphStyle("h2",fontName="Helvetica-Bold",fontSize=11,textColor=_BLUE,spaceBefore=10,spaceAfter=4)
                        _s_body=ParagraphStyle("b",fontName="Helvetica",fontSize=10,textColor=_DARK,spaceAfter=4,leading=14)
                        _s_bull=ParagraphStyle("bl",fontName="Helvetica",fontSize=10,textColor=_DARK,spaceAfter=3,leftIndent=14,leading=13)
                        _logo_path=_os.path.join(_os.path.dirname(_os.path.abspath(__file__)),"logo_viamar.jpg")
                        _elems=[]
                        if _os.path.exists(_logo_path):
                            try:
                                _img=_RLImage(_logo_path,width=5*cm,height=1.6*cm,kind="proportional")
                                _elems.append(_img)
                                _elems.append(Spacer(1,0.3*cm))
                            except Exception: pass
                        _elems+=[Paragraph("NautiCRM — Informe IA",_s_tit),
                                Paragraph(f"Generado: {date.today().strftime('%d/%m/%Y')}",_s_sub),
                                HRFlowable(width="100%",thickness=1,color=_GOLD,spaceAfter=10)]
                        for _ln in txt.split("\n"):
                            _ln=_ln.strip()
                            if not _ln: _elems.append(Spacer(1,0.18*cm)); continue
                            def _bold(_t): return _re.sub(r'\*\*(.*?)\*\*',r'<b>\1</b>',_t)
                            if _ln.startswith("## "): _elems.append(Paragraph(_ln[3:],_s_h1))
                            elif _ln.startswith("### "): _elems.append(Paragraph(_ln[4:],_s_h2))
                            elif _ln.startswith(("- ","* ")): _elems.append(Paragraph("• "+_bold(_ln[2:]),_s_bull))
                            else: _elems.append(Paragraph(_bold(_ln),_s_body))
                        _doc.build(_elems)
                        return _buf.getvalue()

                    def _to_docx(txt):
                        import os as _os
                        from docx import Document
                        from docx.shared import Pt, RGBColor, Cm
                        _doc2=Document(); _doc2.core_properties.author="NautiCRM"
                        _logo_path=_os.path.join(_os.path.dirname(_os.path.abspath(__file__)),"logo_viamar.jpg")
                        if _os.path.exists(_logo_path):
                            try: _doc2.add_picture(_logo_path,width=Cm(5))
                            except Exception: pass
                        _h=_doc2.add_heading("NautiCRM — Informe IA",0)
                        _h.runs[0].font.color.rgb=RGBColor(0xC9,0xA8,0x4C)
                        _doc2.add_paragraph(f"Generado: {date.today().strftime('%d/%m/%Y')}")
                        def _add_run(para,line):
                            for _i,_p in enumerate(_re.split(r'\*\*(.*?)\*\*',line)):
                                _r=para.add_run(_p); _r.bold=(_i%2==1)
                        for _ln in txt.split("\n"):
                            _ln=_ln.strip()
                            if not _ln: _doc2.add_paragraph(""); continue
                            if _ln.startswith("## "): _doc2.add_heading(_ln[3:],level=1)
                            elif _ln.startswith("### "): _doc2.add_heading(_ln[4:],level=2)
                            elif _ln.startswith(("- ","* ")): _add_run(_doc2.add_paragraph(style="List Bullet"),_ln[2:])
                            else: _add_run(_doc2.add_paragraph(),_ln)
                        _buf2=_io_ia.BytesIO(); _doc2.save(_buf2); return _buf2.getvalue()

                    _fname=f"informe_ia_{date.today()}"
                    _dc1,_dc2,_dc3=st.columns(3)
                    _dc1.download_button("⬇️ Descargar TXT",data=_informe.encode("utf-8"),file_name=f"{_fname}.txt",mime="text/plain",use_container_width=True)
                    try:
                        _dc2.download_button("⬇️ Descargar Word",data=_to_docx(_informe),file_name=f"{_fname}.docx",mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",use_container_width=True)
                    except Exception as _ex2:
                        _dc2.warning(f"Word no disponible: {_ex2}")
                    try:
                        _dc3.download_button("⬇️ Descargar PDF",data=_to_pdf(_informe),file_name=f"{_fname}.pdf",mime="application/pdf",use_container_width=True)
                    except Exception as _ex3:
                        _dc3.warning(f"PDF no disponible: {_ex3}")
                except Exception as _e:
                    st.error(f"Error al generar el informe: {_e}")

    st.markdown("---")
    csv_inf=pd.DataFrame([{k:v for k,v in l.items() if k!="historial"} for l in all_leads_raw]).to_csv(index=False).encode("utf-8")
    st.download_button("⬇️ Exportar CSV",data=csv_inf,file_name="nauticrm_informe.csv",mime="text/csv")

# ══ LEAD FORM ═════════════════════════════════════════════════════════════════
elif "Lead" in page:
    st.markdown("## ➕ Nuevo / Editar Lead")

    # ── Confirmación de borrado (prioridad máxima, bloquea el resto) ──────────
    if st.session_state.get("pending_delete_id"):
        _del_id     = st.session_state["pending_delete_id"]
        _del_nombre = st.session_state["pending_delete_nombre"]
        st.warning(f"⚠️ ¿Seguro que quieres eliminar a **{_del_nombre}**? Esta acción no se puede deshacer.")
        _c1, _c2 = st.columns(2)
        if _c1.button("🗑️ Sí, eliminar definitivamente", use_container_width=True):
            delete_lead(_del_id)
            st.session_state["pending_delete_id"]     = None
            st.session_state["pending_delete_nombre"] = None
            st.session_state["_nav_request"] = "⊞ Funnel Kanban"
            st.rerun()
        if _c2.button("↩ Cancelar", use_container_width=True):
            st.session_state["pending_delete_id"]     = None
            st.session_state["pending_delete_nombre"] = None
            st.rerun()
        st.stop()

    # ── Selector de lead ──────────────────────────────────────────────────────
    leads_editables=sorted(all_leads_raw,key=lambda l:l.get("fechaCreacion",""),reverse=True)
    nombres_lista=["— Crear nuevo lead —"]+[l["nombre"] for l in leads_editables]
    _goto=st.session_state.get("goto_lead")
    if _goto:
        st.session_state["sel_lead"] = _goto if _goto in nombres_lista else nombres_lista[0]
        st.session_state["goto_lead"] = None
    if st.session_state.get("_sel_lead_request"):
        st.session_state["sel_lead"] = st.session_state["_sel_lead_request"]
        st.session_state["_sel_lead_request"] = None
    if "sel_lead" not in st.session_state:
        st.session_state["sel_lead"] = nombres_lista[0]
    sel=st.selectbox("Seleccionar lead existente",nombres_lista,key="sel_lead")
    existing=None if sel=="— Crear nuevo lead —" else next((l for l in leads_editables if l["nombre"]==sel),None)
    d=existing or {}
    if existing:
        st.markdown("---")
        st.markdown("##### ⚡ Cambio rápido de estado")
        c1,c2=st.columns([3,1])
        etapa_r=c1.selectbox("Nueva etapa",STAGES,index=STAGES.index(existing["etapa"]) if existing["etapa"] in STAGES else 0,key="etapa_r")
        if c2.button("✅ Actualizar"):
            existing["etapa"]=etapa_r; existing["ultimaActualizacion"]=str(date.today())
            save_lead(existing,is_new=False); st.success(f"✅ → {etapa_r}"); st.rerun()
        st.markdown("---")
    if st.session_state.get("_ultimo_guardado"):
        st.success(f"✅ Lead **{st.session_state['_ultimo_guardado']}** guardado correctamente.")
        st.session_state["_ultimo_guardado"] = None

    with st.form(f"lead_form_{d.get('id','new')}"):
        st.markdown("### 👤 Datos de Contacto")
        c1,c2,c3=st.columns(3)
        nombre =c1.text_input("Nombre *",value=d.get("nombre",""))
        empresa=c2.text_input("Empresa",value=d.get("empresa",""))
        idioma =c3.selectbox("Idioma",IDIOMAS,index=IDIOMAS.index(d["idioma"]) if d.get("idioma") in IDIOMAS else 0)
        c4,c5=st.columns(2)
        tel  =c4.text_input("Teléfono",value=d.get("telefono",""))
        email=c5.text_input("Email",value=d.get("email",""))
        st.markdown("### 🚢 Embarcación de Interés")
        c6,c7=st.columns(2)
        tipo  =c6.selectbox("Tipo / Marca",boat_types,index=boat_types.index(d["tipoEmbarcacion"]) if d.get("tipoEmbarcacion") in boat_types else 0)
        modelo=c7.text_input("Modelo / Eslora",value=d.get("modeloEslora",""))
        uso   =st.text_area("Notas internas",value=d.get("usoPrevisto",""),placeholder="Observaciones internas, contexto del cliente...",height=80)
        st.markdown("### 📋 Gestión Comercial")
        c9,c10,c11,c12=st.columns(4)
        etapa =c9.selectbox("Etapa",STAGES,index=STAGES.index(d["etapa"]) if d.get("etapa") in STAGES else 0)
        asig  =c10.selectbox("Asignado a",vendedores,index=vendedores.index(d["asignadoA"]) if d.get("asignadoA") in vendedores else 0)
        fuente=c11.selectbox("Fuente",sources,index=sources.index(d["fuenteLead"]) if d.get("fuenteLead") in sources else 0)
        prob  =c12.number_input("Probabilidad (%)",value=float(d.get("probabilidad",20) or 20),min_value=0.0,max_value=100.0,step=5.0)
        c13,c14,c15=st.columns(3)
        valor =c13.number_input("Valor operación (€)",value=float(d.get("valorOperacion",0) or 0),min_value=0.0,step=5000.0)
        prox_a=c14.text_input("Próxima acción",value=d.get("proximaAccion",""))
        try:    prox_d=c15.date_input("Fecha próx. acción",value=datetime.strptime(str(d.get("fechaProximaAccion","")),"%Y-%m-%d").date())
        except: prox_d=c15.date_input("Fecha próx. acción",value=date.today())
        submitted=st.form_submit_button("💾 Guardar Lead",use_container_width=True)
        if submitted:
            if not nombre.strip(): st.error("El nombre es obligatorio.")
            elif st.session_state.get("_guardando"):
                st.warning("⏳ Guardando, por favor espera...")
            else:
                st.session_state["_guardando"] = True
                new_lead={"id":d.get("id","") or str(uuid.uuid4()),"nombre":nombre,"empresa":empresa,"telefono":tel,"email":email,"idioma":idioma,"tipoEmbarcacion":tipo,"modeloEslora":modelo,"presupuesto":int(valor),"usoPrevisto":uso,"asignadoA":asig,"etapa":etapa,"probabilidad":int(prob),"valorOperacion":int(valor),"fuenteLead":fuente,"proximaAccion":prox_a,"fechaProximaAccion":str(prox_d),"historial":d.get("historial",[]),"fechaCreacion":d.get("fechaCreacion","") or str(date.today()),"ultimaActualizacion":str(date.today())}
                save_lead(new_lead,is_new=not bool(existing))
                st.session_state["_guardando"] = False
                st.session_state["_ultimo_guardado"] = nombre
                st.session_state["_sel_lead_request"] = nombre
                st.rerun()

    if existing:
        # ── Registrar actividad ───────────────────────────────────────────────
        st.markdown("---")
        st.markdown("##### 📝 Registrar actividad")
        _a1, _a2 = st.columns([1, 3])
        _act_tipo = _a1.selectbox("Tipo de contacto", ["Email","Llamada","Reunión","WhatsApp","Nota"], key="act_tipo")
        _act_nota = _a2.text_area("Descripción", placeholder="Qué se dijo, qué se envió, qué contestó el lead...", height=90, key="act_nota", label_visibility="collapsed")
        st.markdown("<div style='color:#7a8fa6;font-size:0.75rem;margin:6px 0 2px'>↳ Próximo paso (opcional)</div>", unsafe_allow_html=True)
        _b1, _b2 = st.columns([3, 1])
        _prox_a = _b1.text_input("Próxima acción", value=existing.get("proximaAccion",""), placeholder="Ej: Llamada de seguimiento, Enviar contrato...", key="act_prox_a", label_visibility="collapsed")
        try:    _prox_d_def = datetime.strptime(str(existing.get("fechaProximaAccion","")), "%Y-%m-%d").date()
        except: _prox_d_def = date.today() + timedelta(days=7)
        _prox_d = _b2.date_input("Fecha", value=_prox_d_def, key="act_prox_d", label_visibility="collapsed")
        if st.button("➕ Registrar actividad y actualizar plan", use_container_width=True):
            if not _act_nota.strip():
                st.warning("Escribe la descripción antes de registrar.")
            else:
                _hist = existing.get("historial", []).copy()
                _hist.append({"fecha": str(date.today()), "tipo": _act_tipo, "nota": _act_nota.strip()})
                _upd = {**existing, "historial": _hist, "ultimaActualizacion": str(date.today())}
                if _prox_a.strip():
                    _upd["proximaAccion"]      = _prox_a.strip()
                    _upd["fechaProximaAccion"] = str(_prox_d)
                save_lead(_upd, is_new=False)
                st.session_state["_nav_request"] = "📅 Próximas Acciones"
                st.rerun()

        # ── Historial de comunicaciones ───────────────────────────────────────
        st.markdown("---")
        st.markdown("##### 🕐 Historial de comunicaciones")
        _hist_list = existing.get("historial", [])
        if not _hist_list:
            st.caption("Sin actividad registrada aún.")
        else:
            TIPO_COLOR = {"Email":"#2563eb","Llamada":"#16a34a","Reunión":"#7c3aed","WhatsApp":"#25d366","Nota":"#c9a84c"}
            for h in reversed(_hist_list):
                _col = TIPO_COLOR.get(h["tipo"], "#7a8fa6")
                st.markdown(f"""<div style='background:#091220;border:1px solid #1a3050;border-left:3px solid {_col};border-radius:6px;padding:10px 14px;margin-bottom:6px'>
                    <div style='display:flex;align-items:center;gap:8px;margin-bottom:5px'>
                        <span style='color:#c9a84c;font-size:0.78rem;font-weight:700'>{h['fecha']}</span>
                        <span style='background:{_col};color:white;border-radius:4px;padding:1px 8px;font-size:0.68rem;font-weight:700'>{h['tipo']}</span>
                    </div>
                    <div style='color:#e8e0d0;font-size:0.84rem;white-space:pre-wrap'>{h['nota']}</div>
                </div>""", unsafe_allow_html=True)

        st.markdown("---")
        if st.button("🗑️ Eliminar este lead"):
            st.session_state["pending_delete_id"]     = existing["id"]
            st.session_state["pending_delete_nombre"] = existing["nombre"]
            st.rerun()

# ══ PRÓXIMAS ACCIONES ══════════════════════════════════════════════════════════
elif "Acciones" in page:
    st.markdown("## 📅 Próximas Acciones")
    c1,c2,c3=st.columns(3)
    fv_pa=c1.selectbox("Vendedor",["Todos"]+vendedores,key="fv_pa")
    fe_pa=c2.selectbox("Etapa",["Todas"]+FUNNEL_STAGES,key="fe_pa")
    solo_v=c3.checkbox("Solo vencidas / urgentes")
    acciones=[l for l in all_leads_raw if (l.get("proximaAccion") or l.get("fechaProximaAccion")) and l["etapa"] not in ["Cerrado Ganado","Cerrado Perdido"]]
    if fv_pa!="Todos": acciones=[l for l in acciones if l["asignadoA"]==fv_pa]
    if fe_pa!="Todas": acciones=[l for l in acciones if l["etapa"]==fe_pa]
    if solo_v: acciones=[l for l in acciones if urg_emoji(l.get("fechaProximaAccion")) in ["🔴","🟡"]]
    acciones.sort(key=lambda l:(lambda ds:(datetime.strptime(str(ds),"%Y-%m-%d").date() if ds else date(9999,12,31)))(l.get("fechaProximaAccion","")))
    sin_accion=[l for l in all_leads_raw if not l.get("proximaAccion") and not l.get("fechaProximaAccion") and l["etapa"] not in ["Cerrado Ganado","Cerrado Perdido","En Pausa / Recuperable"]]
    k1,k2,k3,k4=st.columns(4)
    k1.metric("🔴 Vencidas",len([l for l in acciones if urg_emoji(l.get("fechaProximaAccion"))=="🔴"]))
    k2.metric("🟡 Hoy",len([l for l in acciones if urg_emoji(l.get("fechaProximaAccion"))=="🟡"]))
    k3.metric("🟢 Próximas",len([l for l in acciones if urg_emoji(l.get("fechaProximaAccion"))=="🟢"]))
    k4.metric("⚪ Sin acción",len(sin_accion))
    st.markdown("<br>", unsafe_allow_html=True)
    current_sec=None
    for l in acciones:
        try:
            diff=(datetime.strptime(str(l.get("fechaProximaAccion","")),"%Y-%m-%d").date()-date.today()).days
            sec="🔴 Vencidas" if diff<0 else "🟡 Para hoy" if diff==0 else "📆 Esta semana" if diff<=7 else "🟢 Próximas"
            dias_txt=f"Venció hace {abs(diff)}d" if diff<0 else "Hoy" if diff==0 else f"En {diff}d"
            dtcol="#e74c3c" if diff<0 else "#f1c40f" if diff==0 else "#2ecc71"
        except: sec,dias_txt,dtcol="⚪ Sin fecha","Sin fecha","#7a8fa6"
        if sec!=current_sec: current_sec=sec; st.markdown(f"### {sec}")
        color=STAGE_COLORS.get(l["etapa"],"#1a3050")
        _pc1,_pc2=st.columns([11,1])
        with _pc1:
            st.markdown(f"""<div style="background:#0d1e35;border:1px solid #1a3050;border-left:4px solid {color};border-radius:8px;padding:12px 16px;margin-bottom:2px;display:flex;justify-content:space-between;align-items:center">
                <div style="flex:1"><div style="display:flex;align-items:center;gap:8px;margin-bottom:4px;flex-wrap:wrap">
                    <span style="color:#e8e0d0;font-weight:700;font-size:0.88rem">{l['nombre']}</span>
                    <span style="background:{color};color:white;border-radius:10px;padding:1px 8px;font-size:0.62rem;font-weight:700">{l['etapa']}</span>
                    <span style="color:#7a8fa6;font-size:0.72rem">{l['asignadoA'].replace('Vendedor','V.')}</span></div>
                <div style="color:#c9a84c;font-size:0.82rem;margin-bottom:3px">↳ {l.get('proximaAccion','Sin descripción')}</div>
                <div style="color:#7a8fa6;font-size:0.72rem">{l.get('empresa','')} · {l.get('tipoEmbarcacion','')} {l.get('modeloEslora','')}</div></div>
                <div style="text-align:right;flex-shrink:0;margin-left:16px">
                    <div style="color:#e8e0d0;font-size:0.82rem;font-weight:700">{l.get('fechaProximaAccion','—')}</div>
                    <div style="font-size:0.72rem;color:{dtcol}">{dias_txt}</div>
                    <div style="color:#2ecc71;font-family:monospace;font-size:0.72rem;margin-top:2px">{fmt_eur(l.get('valorOperacion',0))}</div>
                </div></div>""", unsafe_allow_html=True)
        with _pc2:
            if st.button("✏️", key=f"pa_{l['id']}", help="Ir a ficha y registrar actividad"):
                st.session_state["goto_lead"] = l["nombre"]
                st.session_state["_nav_request"] = "➕ Nuevo / Editar Lead"
                st.rerun()
    if sin_accion:
        st.markdown("---"); st.markdown("### ⚪ Sin próxima acción")
        for l in sin_accion:
            color=STAGE_COLORS.get(l["etapa"],"#1a3050")
            st.markdown(f"<div style='background:#0d1e35;border:1px solid #1a3050;border-left:4px solid {color};border-radius:8px;padding:10px 16px;margin-bottom:6px'><span style='color:#e8e0d0;font-weight:600'>{l['nombre']}</span><span style='background:{color};color:white;border-radius:10px;padding:1px 8px;font-size:0.62rem;font-weight:700;margin:0 8px'>{l['etapa']}</span><span style='color:#7a8fa6;font-size:0.75rem'>{l.get('empresa','')}·{l['asignadoA'].replace('Vendedor','V.')}</span></div>", unsafe_allow_html=True)

    # ── PDF Próximas Acciones ──────────────────────────────────────────────
    st.markdown("---")
    if st.button("🖨️ Generar PDF — Agenda de hoy y esta semana"):
        from reportlab.lib.pagesizes import A4, landscape
        from reportlab.lib import colors
        from reportlab.lib.units import cm
        from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.enums import TA_LEFT, TA_CENTER
        import io as _io_pdf
        _buf_pdf=_io_pdf.BytesIO()
        doc=SimpleDocTemplate(_buf_pdf,pagesize=landscape(A4),leftMargin=1.5*cm,rightMargin=1.5*cm,topMargin=1.5*cm,bottomMargin=1.5*cm)
        styles=getSampleStyleSheet()
        NAVY_C=colors.HexColor("#0A1628"); GOLD_C=colors.HexColor("#C9A84C"); LIGHT_C=colors.HexColor("#E8E0D0"); MID_C=colors.HexColor("#0D1E35")
        RED_C=colors.HexColor("#DC2626"); YEL_C=colors.HexColor("#B8860B"); GRN_C=colors.HexColor("#16A34A"); BLU_C=colors.HexColor("#2563EB")
        title_style=ParagraphStyle("title",fontName="Helvetica-Bold",fontSize=16,textColor=GOLD_C,spaceAfter=4)
        sub_style=ParagraphStyle("sub",fontName="Helvetica",fontSize=9,textColor=LIGHT_C,spaceAfter=12)
        sec_style=ParagraphStyle("sec",fontName="Helvetica-Bold",fontSize=11,textColor=GOLD_C,spaceBefore=12,spaceAfter=6)
        cell_style=ParagraphStyle("cell",fontName="Helvetica",fontSize=8,textColor=colors.black,leading=10)
        elems=[]
        elems.append(Paragraph("⚓  NautiCRM — Agenda de Próximas Acciones", title_style))
        elems.append(Paragraph(f"Generado: {date.today().strftime('%d/%m/%Y')}  ·  Vendedor: {fv_pa if fv_pa!='Todos' else 'Todos'}", sub_style))
        grupos=[("🔴 Vencidas", [l for l in acciones if urg_emoji(l.get("fechaProximaAccion"))=="🔴"], RED_C),
                ("🟡 Para hoy",  [l for l in acciones if urg_emoji(l.get("fechaProximaAccion"))=="🟡"], YEL_C),
                ("📆 Esta semana",[l for l in acciones if urg_emoji(l.get("fechaProximaAccion"))=="🟢" and (datetime.strptime(str(l.get("fechaProximaAccion","")),"%Y-%m-%d").date()-date.today()).days<=7], BLU_C),
                ("🟢 Próximas",  [l for l in acciones if urg_emoji(l.get("fechaProximaAccion"))=="🟢" and (datetime.strptime(str(l.get("fechaProximaAccion","")),"%Y-%m-%d").date()-date.today()).days>7], GRN_C),
        ]
        col_w=[5*cm,5*cm,2.5*cm,7*cm,2.5*cm,4*cm]
        for sec_name,sec_leads,sec_color in grupos:
            if not sec_leads: continue
            elems.append(Paragraph(sec_name, sec_style))
            tdata=[["Cliente","Empresa","Fecha","Próxima Acción","Etapa","Valor"]]
            for l in sec_leads:
                tdata.append([
                    Paragraph(l.get("nombre",""),cell_style),
                    Paragraph(l.get("empresa",""),cell_style),
                    l.get("fechaProximaAccion","—"),
                    Paragraph(l.get("proximaAccion",""),cell_style),
                    Paragraph(l.get("etapa",""),cell_style),
                    fmt_eur(l.get("valorOperacion",0)),
                ])
            t=Table(tdata,colWidths=col_w,repeatRows=1)
            t.setStyle(TableStyle([
                ("BACKGROUND",(0,0),(-1,0),NAVY_C),("TEXTCOLOR",(0,0),(-1,0),GOLD_C),
                ("FONTNAME",(0,0),(-1,0),"Helvetica-Bold"),("FONTSIZE",(0,0),(-1,0),8),
                ("FONTNAME",(0,1),(-1,-1),"Helvetica"),("FONTSIZE",(0,1),(-1,-1),8),
                ("BACKGROUND",(0,1),(-1,-1),colors.HexColor("#F5F5F5")),
                ("ROWBACKGROUNDS",(0,1),(-1,-1),[colors.white,colors.HexColor("#F0F4F8")]),
                ("GRID",(0,0),(-1,-1),0.5,colors.HexColor("#CCCCCC")),
                ("LEFTPADDING",(0,0),(-1,-1),4),("RIGHTPADDING",(0,0),(-1,-1),4),
                ("TOPPADDING",(0,0),(-1,-1),4),("BOTTOMPADDING",(0,0),(-1,-1),4),
                ("VALIGN",(0,0),(-1,-1),"TOP"),
                ("LINEBELOW",(0,0),(-1,0),1.5,sec_color),
            ]))
            elems.append(t)
            elems.append(Spacer(1,0.3*cm))
        doc.build(elems)
        st.download_button("⬇️ Descargar PDF",data=_buf_pdf.getvalue(),file_name=f"agenda_acciones_{date.today()}.pdf",mime="application/pdf")

# ══ ARCHIVO FRÍO ══════════════════════════════════════════════════════════════
elif "Archivo" in page:
    st.markdown("## 🧊 Archivo Frío")
    st.markdown("Contactos en pausa más de **6 meses**.")
    pausados=[l for l in all_leads_raw if l.get("etapa")=="En Pausa / Recuperable"]
    if pausados:
        with st.expander(f"📥 Archivar manualmente ({len(pausados)} en pausa)"):
            for l in pausados:
                c1,c2=st.columns([4,1])
                c1.markdown(f"**{l['nombre']}** — {l['empresa']} · {months_since(l.get('ultimaActualizacion') or l.get('fechaCreacion',''))} meses")
                if c2.button("Archivar",key=f"arch_{l['id']}"):
                    archivo_frio.append({**l,"fechaArchivo":str(date.today())})
                    save_config(vendedores,boat_types,sources,archivo_frio)
                    delete_lead(l["id"]); st.rerun()
    if not archivo_frio: st.info("El archivo frío está vacío.")
    else:
        st.markdown("#### 🔎 Segmentación")
        _fa1,_fa2,_fa3=st.columns(3)
        q_a   =_fa1.text_input("🔍 Buscar nombre / empresa",key="q_arch")
        ft_a  =_fa2.selectbox("Embarcación",["Todos"]+boat_types,key="ft_arch")
        fi_a  =_fa3.selectbox("Idioma",["Todos"]+IDIOMAS,key="fi_arch")
        _fb1,_fb2,_fb3=st.columns(3)
        fv_a  =_fb1.selectbox("Vendedor",["Todos"]+vendedores,key="fv_arch")
        _presu_vals=[l.get("valorOperacion",0) or l.get("presupuesto",0) for l in archivo_frio if l.get("valorOperacion",0) or l.get("presupuesto",0)]
        _pmin=int(min(_presu_vals)) if _presu_vals else 0
        _pmax=int(max(_presu_vals)) if _presu_vals else 500000
        if _pmax>_pmin:
            _rango=_fb2.slider("Presupuesto (€)",min_value=_pmin,max_value=_pmax,value=(_pmin,_pmax),step=5000,key="fp_arch",format="€%d")
        else:
            _rango=(_pmin,_pmax)

        filtrado=archivo_frio
        if q_a:        filtrado=[l for l in filtrado if q_a.lower() in (l.get("nombre","")+" "+l.get("empresa","")).lower()]
        if ft_a!="Todos": filtrado=[l for l in filtrado if l.get("tipoEmbarcacion")==ft_a]
        if fi_a!="Todos": filtrado=[l for l in filtrado if l.get("idioma")==fi_a]
        if fv_a!="Todos": filtrado=[l for l in filtrado if l.get("asignadoA")==fv_a]
        filtrado=[l for l in filtrado if _rango[0]<=(l.get("valorOperacion",0) or l.get("presupuesto",0))<=_rango[1]]

        st.caption(f"**{len(filtrado)}** contactos seleccionados de {len(archivo_frio)} en total")
        rows=[{
            "Nombre":l.get("nombre",""),"Empresa":l.get("empresa",""),
            "Idioma":l.get("idioma","—"),
            "Embarcación":f"{l.get('tipoEmbarcacion','')} {l.get('modeloEslora','')}".strip(),
            "Valor €":fmt_eur(l.get("valorOperacion",0) or l.get("presupuesto",0)),
            "Email":l.get("email",""),"Teléfono":l.get("telefono",""),
            "Vendedor":l.get("asignadoA",""),"Archivado":l.get("fechaArchivo","")
        } for l in filtrado]
        st.dataframe(pd.DataFrame(rows),use_container_width=True,hide_index=True)

        import io as _io2
        rows_xl=[{
            "Nombre":l.get("nombre",""),"Empresa":l.get("empresa",""),
            "Email":l.get("email",""),"Teléfono":l.get("telefono",""),
            "Idioma":l.get("idioma","—"),
            "Embarcación":f"{l.get('tipoEmbarcacion','')} {l.get('modeloEslora','')}".strip(),
            "Notas internas":l.get("usoPrevisto",""),
            "Valor €":l.get("valorOperacion",0) or l.get("presupuesto",0),
            "Vendedor":l.get("asignadoA",""),
            "Fecha entrada":l.get("fechaCreacion",""),
            "Archivado":l.get("fechaArchivo","")
        } for l in filtrado]
        _sufijo="_".join(filter(None,[ft_a if ft_a!="Todos" else "",fi_a if fi_a!="Todos" else "",fv_a if fv_a!="Todos" else ""]))
        _fname=f"archivo_frio{'_'+_sufijo if _sufijo else ''}_{date.today()}.xlsx"
        _buf_arch=_io2.BytesIO()
        pd.DataFrame(rows_xl).to_excel(_buf_arch,index=False,engine="openpyxl")
        st.download_button("⬇️ Exportar segmento a Excel",data=_buf_arch.getvalue(),file_name=_fname,mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
        st.markdown("---"); st.markdown("**♻️ Reactivar contacto**")
        sel_r=st.selectbox("Seleccionar",["— Selecciona —"]+[l["nombre"] for l in archivo_frio])
        if sel_r!="— Selecciona —" and st.button("♻️ Reactivar como Prospecto"):
            lr=next(l for l in archivo_frio if l["nombre"]==sel_r)
            nuevo_lead={**lr,"etapa":"Prospecto","ultimaActualizacion":str(date.today())}
            nuevo_lead.pop("fechaArchivo",None)
            save_lead(nuevo_lead,is_new=True)
            archivo_frio_nuevo=[l for l in archivo_frio if l["nombre"]!=sel_r]
            save_config(vendedores,boat_types,sources,archivo_frio_nuevo)
            st.success(f"✅ {sel_r} reactivado."); st.rerun()

# ══ CONFIG ════════════════════════════════════════════════════════════════════
elif "Config" in page:
    st.markdown("## ⚙️ Configuración")
    tab1,tab2,tab3=st.tabs(["👥 Equipo de ventas","🚢 Catálogo de embarcaciones","📡 Fuentes de lead"])
    with tab1:
        with st.form("cfg_v"):
            names=[st.text_input(f"Vendedor {i+1}",value=v) for i,v in enumerate(vendedores)]
            if st.form_submit_button("💾 Guardar nombres"):
                save_config(names,boat_types,sources,archivo_frio); st.success("✅ Actualizado."); st.rerun()
    with tab2:
        st.caption("Gestiona los valores del desplegable de embarcación.")
        st.markdown("<br>", unsafe_allow_html=True)
        cols_bt=st.columns(4); to_delete=None
        for i,bt in enumerate(boat_types):
            with cols_bt[i%4]:
                c1,c2=st.columns([3,1])
                c1.markdown(f"<div style='background:#0d1e35;border:1px solid #1a3050;border-radius:6px;padding:5px 10px;font-size:0.82rem;color:#e8e0d0'>{bt}</div>", unsafe_allow_html=True)
                if c2.button("✕",key=f"del_{i}"): to_delete=bt
        if to_delete:
            new_bt=[x for x in boat_types if x!=to_delete]
            save_config(vendedores,new_bt,sources,archivo_frio); st.rerun()
        st.markdown("<br>", unsafe_allow_html=True)
        with st.form("cfg_b"):
            c1,c2=st.columns([3,1])
            nueva=c1.text_input("Nueva marca/tipo",placeholder="Ej: Hallberg-Rassy, Dufour...",label_visibility="collapsed")
            if c2.form_submit_button("➕ Añadir"):
                if nueva.strip() and nueva.strip() not in boat_types:
                    save_config(vendedores,boat_types+[nueva.strip()],sources,archivo_frio); st.success(f"✅ '{nueva.strip()}' añadido."); st.rerun()
                elif nueva.strip() in boat_types: st.warning("Ya existe.")
        if st.button("↺ Restaurar por defecto"):
            save_config(vendedores,["Velero","Motor","Catamarán","Zodiac","Charter","Jeanneau","Beneteau","Sunseeker","Princess","Azimut","Ferretti","Bavaria","Hanse","Lagoon","Otro"],sources,archivo_frio); st.rerun()
    with tab3:
        st.caption("Gestiona los valores del desplegable de fuente de lead.")
        st.markdown("<br>", unsafe_allow_html=True)
        cols_src=st.columns(4); to_delete_src=None
        for i,src in enumerate(sources):
            with cols_src[i%4]:
                c1,c2=st.columns([3,1])
                c1.markdown(f"<div style='background:#0d1e35;border:1px solid #1a3050;border-radius:6px;padding:5px 10px;font-size:0.82rem;color:#e8e0d0'>{src}</div>", unsafe_allow_html=True)
                if c2.button("✕",key=f"del_src_{i}"): to_delete_src=src
        if to_delete_src:
            new_src=[x for x in sources if x!=to_delete_src]
            save_config(vendedores,boat_types,new_src,archivo_frio); st.rerun()
        st.markdown("<br>", unsafe_allow_html=True)
        with st.form("cfg_src"):
            c1,c2=st.columns([3,1])
            nueva_src=c1.text_input("Nueva fuente",placeholder="Ej: LinkedIn, Partner, Evento...",label_visibility="collapsed")
            if c2.form_submit_button("➕ Añadir"):
                if nueva_src.strip() and nueva_src.strip() not in sources:
                    save_config(vendedores,boat_types,sources+[nueva_src.strip()],archivo_frio); st.success(f"✅ '{nueva_src.strip()}' añadido."); st.rerun()
                elif nueva_src.strip() in sources: st.warning("Ya existe.")
        if st.button("↺ Restaurar fuentes por defecto"):
            save_config(vendedores,boat_types,["Feria Náutica","Web","Referido","RRSS","Llamada Fría","Otro"],archivo_frio); st.rerun()
    st.markdown("---")
    st.markdown("##### 📱 Enlace Modo Feria")
    app_url=st.secrets.get("app_url","https://nauticrm-trn2jrtqldn9vfd4zic6hz.streamlit.app")
    st.code(f"{app_url}/?modo=feria")
    st.caption("Comparte esta URL con los vendedores para registro rápido desde móvil en ferias.")
